from __future__ import annotations

import html
import os
import re
from datetime import datetime, timedelta, timezone
from pathlib import Path
from urllib.parse import quote
from uuid import UUID

import anyio
from aiofile import async_open
from loguru import logger

from agentcore.services.settings.service import SettingsService

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
OFFICE_VIEWER_BASE_URL = "https://view.officeapps.live.com/op/embed.aspx?src="


def sanitize_release_document_name(file_name: str) -> str:
    cleaned = Path(file_name or "release-notes.docx").name.strip()
    if not cleaned:
        cleaned = "release-notes.docx"
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", cleaned)
    return cleaned or "release-notes.docx"


def build_release_document_path(release_id: UUID | str, file_name: str) -> str:
    return f"releases/{release_id}/{sanitize_release_document_name(file_name)}"


def _get_release_documents_container(settings_service: SettingsService) -> str:
    configured = str(
        getattr(settings_service.settings, "azure_release_documents_container_name", "") or ""
    ).strip()
    if not configured:
        raise ValueError(
            "AZURE_RELEASE_DOCUMENTS_CONTAINER_NAME is required for release document storage."
        )
    return configured


def _get_storage_account_url() -> str:
    account_url = os.environ.get("AZURE_STORAGE_ACCOUNT_URL", "").strip().strip("'\"")
    if not account_url:
        raise ValueError("AZURE_STORAGE_ACCOUNT_URL is required when STORAGE_TYPE=azure.")
    return account_url.rstrip("/")


async def save_release_document(
    *,
    settings_service: SettingsService,
    release_id: UUID | str,
    file_name: str,
    content: bytes,
) -> str:
    storage_type = str(getattr(settings_service.settings, "storage_type", "local") or "local").strip().lower()
    blob_path = build_release_document_path(release_id, file_name)

    if storage_type == "azure":
        from azure.identity.aio import DefaultAzureCredential
        from azure.storage.blob.aio import BlobServiceClient

        account_url = _get_storage_account_url()
        container_name = _get_release_documents_container(settings_service)
        credential = DefaultAzureCredential(
            exclude_environment_credential=True,
            exclude_interactive_browser_credential=True,
        )
        blob_service_client = BlobServiceClient(account_url=account_url, credential=credential)
        container_client = blob_service_client.get_container_client(container_name)
        try:
            try:
                await container_client.get_container_properties()
            except Exception:
                await container_client.create_container()
                logger.info(f"Created Azure blob container: {container_name}")
            await container_client.upload_blob(name=blob_path, data=content, overwrite=True)
        finally:
            await blob_service_client.close()
            close_credential = getattr(credential, "close", None)
            if callable(close_credential):
                await close_credential()
        return blob_path

    base_dir = anyio.Path(settings_service.settings.config_dir) / "release_documents"
    file_path = base_dir / blob_path
    await file_path.parent.mkdir(parents=True, exist_ok=True)
    async with async_open(str(file_path), "wb") as file_handle:
        await file_handle.write(content)
    return blob_path


async def get_release_document(
    *,
    settings_service: SettingsService,
    storage_path: str,
) -> bytes:
    storage_type = str(getattr(settings_service.settings, "storage_type", "local") or "local").strip().lower()

    if storage_type == "azure":
        from azure.identity.aio import DefaultAzureCredential
        from azure.storage.blob.aio import BlobServiceClient

        account_url = _get_storage_account_url()
        container_name = _get_release_documents_container(settings_service)
        credential = DefaultAzureCredential(
            exclude_environment_credential=True,
            exclude_interactive_browser_credential=True,
        )
        blob_service_client = BlobServiceClient(account_url=account_url, credential=credential)
        container_client = blob_service_client.get_container_client(container_name)
        blob_client = container_client.get_blob_client(storage_path)
        try:
            stream = await blob_client.download_blob()
            return await stream.readall()
        finally:
            await blob_service_client.close()
            close_credential = getattr(credential, "close", None)
            if callable(close_credential):
                await close_credential()

    file_path = anyio.Path(settings_service.settings.config_dir) / "release_documents" / storage_path
    if not await file_path.exists():
        raise FileNotFoundError(storage_path)
    async with async_open(str(file_path), "rb") as file_handle:
        return await file_handle.read()


async def build_release_document_office_viewer_url(
    *,
    settings_service: SettingsService,
    storage_path: str,
) -> str | None:
    storage_type = str(getattr(settings_service.settings, "storage_type", "local") or "local").strip().lower()
    if storage_type != "azure":
        return None

    account_url = os.environ.get("AZURE_STORAGE_ACCOUNT_URL", "").strip().strip("'\"")
    if not account_url:
        return None

    from azure.identity.aio import DefaultAzureCredential
    from azure.storage.blob import BlobSasPermissions, generate_blob_sas
    from azure.storage.blob.aio import BlobServiceClient

    container_name = _get_release_documents_container(settings_service)
    credential = DefaultAzureCredential(
        exclude_environment_credential=True,
        exclude_interactive_browser_credential=True,
    )
    blob_service_client = BlobServiceClient(account_url=account_url, credential=credential)
    try:
        now = datetime.now(timezone.utc)
        key_start = now - timedelta(minutes=5)
        key_expiry = now + timedelta(minutes=30)
        delegation_key = await blob_service_client.get_user_delegation_key(
            key_start_time=key_start,
            key_expiry_time=key_expiry,
        )
        account_name = getattr(blob_service_client, "account_name", "") or ""
        if not account_name:
            return None
        sas_token = generate_blob_sas(
            account_name=account_name,
            container_name=container_name,
            blob_name=storage_path,
            user_delegation_key=delegation_key,
            permission=BlobSasPermissions(read=True),
            start=key_start,
            expiry=key_expiry,
        )
        if not sas_token:
            return None
        direct_url = f"{account_url.rstrip('/')}/{container_name}/{quote(storage_path)}?{sas_token}"
        return f"{OFFICE_VIEWER_BASE_URL}{quote(direct_url, safe='')}"
    except Exception as exc:
        logger.warning(f"Failed to build Office viewer URL for release document: {exc}")
        return None
    finally:
        await blob_service_client.close()
        close_credential = getattr(credential, "close", None)
        if callable(close_credential):
            await close_credential()


def render_release_document_preview_html(document_bytes: bytes) -> str:
    from docx import Document
    from io import BytesIO

    document = Document(BytesIO(document_bytes))
    blocks: list[str] = []
    list_items: list[str] = []

    def flush_list_items() -> None:
        nonlocal list_items
        if list_items:
            blocks.append(f"<ul>{''.join(list_items)}</ul>")
            list_items = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            flush_list_items()
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        escaped = html.escape(text)
        if style_name.startswith("heading 1"):
            flush_list_items()
            blocks.append(f"<h1>{escaped}</h1>")
        elif style_name.startswith("heading 2"):
            flush_list_items()
            blocks.append(f"<h2>{escaped}</h2>")
        elif style_name.startswith("heading 3"):
            flush_list_items()
            blocks.append(f"<h3>{escaped}</h3>")
        elif style_name.startswith("list"):
            list_items.append(f"<li>{escaped}</li>")
        else:
            flush_list_items()
            blocks.append(f"<p>{escaped}</p>")

    flush_list_items()

    tables_html: list[str] = []
    for table in document.tables:
        rows_html: list[str] = []
        for row in table.rows:
            cols = "".join(f"<td>{html.escape(cell.text.strip())}</td>" for cell in row.cells)
            rows_html.append(f"<tr>{cols}</tr>")
        tables_html.append(f"<table>{''.join(rows_html)}</table>")

    body = "".join(blocks + tables_html)
    if not body:
        body = "<p>No previewable content found in the uploaded release document.</p>"

    return f"""
    <div class="release-doc-preview">
      {body}
    </div>
    """
